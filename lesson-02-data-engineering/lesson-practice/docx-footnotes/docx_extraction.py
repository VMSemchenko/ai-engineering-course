from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

DOCX_PATH = Path(__file__).parent / "medical_report.docx"

doc = Document(str(DOCX_PATH))


# Step 1: Naive extraction (just paragraphs — misses tracked changes)
print("=" * 60)
print("NAIVE: doc.paragraphs (як більшість парсерів)")
print("=" * 60)

for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"  [{i}] {p.text}")

print("\n-> Tracked changes (insertions/deletions) загублені або змішані!")


# Step 2: Extract tracked changes from XML
print(f"\n{'=' * 60}")
print("TRACKED CHANGES (insertions / deletions)")
print("=" * 60)

insertions = []
deletions = []

for p in doc.paragraphs:
    for ins in p._p.findall(qn("w:ins")):
        for run in ins.findall(qn("w:r")):
            for t in run.findall(qn("w:t")):
                if t.text:
                    author = ins.get(qn("w:author"), "unknown")
                    date = ins.get(qn("w:date"), "")
                    insertions.append({
                        "text": t.text,
                        "author": author,
                        "date": date,
                    })

    for delete in p._p.findall(qn("w:del")):
        for run in delete.findall(qn("w:r")):
            for dt in run.findall(qn("w:delText")):
                if dt.text:
                    author = delete.get(qn("w:author"), "unknown")
                    date = delete.get(qn("w:date"), "")
                    deletions.append({
                        "text": dt.text,
                        "author": author,
                        "date": date,
                    })

print(f"\nInsertions ({len(insertions)}):")
for ins in insertions:
    print(f"  + '{ins['text']}' by {ins['author']} ({ins['date'][:10]})")

print(f"\nDeletions ({len(deletions)}):")
for d in deletions:
    print(f"  - '{d['text']}' by {d['author']} ({d['date'][:10]})")


# Step 3: Extract footnotes from XML
print(f"\n{'=' * 60}")
print("FOOTNOTES")
print("=" * 60)

footnotes_part = doc.part._element.findall(qn("w:footnotes"))
footnotes = []

for fn_container in footnotes_part:
    for fn in fn_container.findall(qn("w:footnote")):
        fn_id = fn.get(qn("w:id"))
        texts = []
        for p in fn.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    if t.text:
                        texts.append(t.text)
        if texts:
            footnotes.append({"id": fn_id, "text": " ".join(texts)})

if footnotes:
    for fn in footnotes:
        print(f"  [{fn['id']}] {fn['text']}")
else:
    print("  (no footnotes in XML — python-docx has limited footnote write support)")
    print("  In real DOCX files from Word, footnotes would appear here.")


# Step 4: Full reconstruction
print(f"\n{'=' * 60}")
print("FULL TEXT (з tracked changes)")
print("=" * 60)

for p in doc.paragraphs:
    parts = []
    for child in p._p:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "r":
            for t in child.findall(qn("w:t")):
                if t.text:
                    parts.append(t.text)

        elif tag == "ins":
            for run in child.findall(qn("w:r")):
                for t in run.findall(qn("w:t")):
                    if t.text:
                        parts.append(f"[+{t.text}]")

        elif tag == "del":
            for run in child.findall(qn("w:r")):
                for dt in run.findall(qn("w:delText")):
                    if dt.text:
                        parts.append(f"[-{dt.text}]")

    line = "".join(parts)
    if line.strip():
        print(f"  {line}")

print("\n-> [+text] = insertion, [-text] = deletion")
print("   Для юридичних/медичних документів tracked changes — критична інформація.")
